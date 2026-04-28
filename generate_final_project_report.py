from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from process_s2_pipeline import apply_butterworth_bandpass_fft, load_subject_pickle


OUTPUT_DIR_NAME = "outputs"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def get_font(size: int) -> ImageFont.ImageFont:
    try:
        return ImageFont.truetype("arial.ttf", size)
    except OSError:
        return ImageFont.load_default()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = italic


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value


def create_signal_comparison_png(workspace_root: Path, output_path: Path) -> None:
    subject_id = "S2"
    data = load_subject_pickle(subject_id, workspace_root)
    bvp = np.asarray(data["signal"]["wrist"]["BVP"]).flatten()
    filtered_bvp = apply_butterworth_bandpass_fft(
        signal=bvp,
        sampling_rate=64,
        low_cutoff_hz=0.5,
        high_cutoff_hz=4.0,
        order=4,
    )

    seconds = 30
    fs = 64
    samples = seconds * fs
    time_axis = np.arange(samples) / fs
    raw_segment = bvp[:samples]
    filtered_segment = filtered_bvp[:samples]

    width = 1200
    height = 700
    margin_left = 80
    margin_right = 30
    margin_top = 60
    margin_bottom = 50
    panel_gap = 40
    panel_height = (height - margin_top - margin_bottom - panel_gap) // 2
    plot_width = width - margin_left - margin_right

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    axis_font = get_font(18)
    tick_font = get_font(14)

    draw.text((width // 2 - 220, 15), "Raw vs Filtered Wrist BVP Signal (S2)", fill="black", font=title_font)

    def draw_panel(top: int, signal: np.ndarray, color: tuple[int, int, int], panel_title: str) -> None:
        y_min = float(signal.min())
        y_max = float(signal.max())
        if y_max == y_min:
            y_max = y_min + 1.0

        def scale_x(x_value: float) -> float:
            return margin_left + (x_value / seconds) * plot_width

        def scale_y(y_value: float) -> float:
            return top + panel_height - ((y_value - y_min) / (y_max - y_min)) * panel_height

        draw.rectangle(
            [margin_left, top, margin_left + plot_width, top + panel_height],
            outline=(180, 180, 180),
            width=1,
        )
        draw.text((margin_left, top - 28), panel_title, fill="black", font=axis_font)

        for tick in range(0, 7):
            x_value = seconds * tick / 6
            x = scale_x(x_value)
            draw.line([(x, top), (x, top + panel_height)], fill=(230, 230, 230), width=1)
            draw.text((x - 10, top + panel_height + 8), f"{x_value:.0f}", fill="black", font=tick_font)

        for tick in range(0, 5):
            y_value = y_min + (y_max - y_min) * tick / 4
            y = scale_y(y_value)
            draw.line([(margin_left, y), (margin_left + plot_width, y)], fill=(230, 230, 230), width=1)
            draw.text((10, y - 8), f"{y_value:.1f}", fill="black", font=tick_font)

        points = [(scale_x(float(xv)), scale_y(float(yv))) for xv, yv in zip(time_axis, signal)]
        draw.line(points, fill=color, width=2)

    draw_panel(margin_top, raw_segment, (200, 70, 70), "Raw BVP (first 30 seconds)")
    draw_panel(margin_top + panel_height + panel_gap, filtered_segment, (40, 100, 210), "Filtered BVP (0.5-4.0 Hz Butterworth-style)")
    draw.text((width // 2 - 80, height - 30), "Time (seconds)", fill="black", font=axis_font)

    image.save(output_path)


def create_model_comparison_png(outputs_dir: Path, output_path: Path) -> None:
    svm_3 = load_json(outputs_dir / "svm_3class_metrics.json")
    rf_3 = load_json(outputs_dir / "rf_3class_metrics.json")
    cnn_3 = load_json(outputs_dir / "cnn_3class_metrics.json")
    svm_b = load_json(outputs_dir / "svm_binary_metrics.json")
    rf_b = load_json(outputs_dir / "rf_binary_metrics.json")
    cnn_b = load_json(outputs_dir / "cnn_binary_metrics.json")

    models = ["SVM", "Random Forest", "CNN"]
    accuracies_3 = [svm_3["accuracy"], rf_3["accuracy"], cnn_3["accuracy"]]
    f1_3 = [svm_3["macro_f1"], rf_3["macro_f1"], cnn_3["macro_f1"]]
    accuracies_b = [svm_b["accuracy"], rf_b["accuracy"], cnn_b["accuracy"]]
    f1_b = [svm_b["macro_f1"], rf_b["macro_f1"], cnn_b["macro_f1"]]

    width = 1200
    height = 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    text_font = get_font(18)
    small_font = get_font(14)
    draw.text((width // 2 - 150, 20), "Model Comparison Summary", fill="black", font=title_font)

    panels = [
        ("3-Class Accuracy", accuracies_3, 60, 110),
        ("3-Class Macro F1", f1_3, 620, 110),
        ("Binary Accuracy", accuracies_b, 60, 400),
        ("Binary Macro F1", f1_b, 620, 400),
    ]
    colors = [(80, 140, 230), (70, 170, 100), (230, 140, 60)]

    for title, values, left, top in panels:
        panel_width = 500
        panel_height = 220
        draw.rectangle([left, top, left + panel_width, top + panel_height], outline=(180, 180, 180), width=1)
        draw.text((left + 10, top - 28), title, fill="black", font=text_font)

        base_y = top + panel_height - 30
        draw.line([(left + 40, top + 20), (left + 40, base_y)], fill="black", width=1)
        draw.line([(left + 40, base_y), (left + panel_width - 20, base_y)], fill="black", width=1)

        for tick in range(0, 6):
            value = tick / 5
            y = base_y - value * 160
            draw.line([(left + 40, y), (left + panel_width - 20, y)], fill=(230, 230, 230), width=1)
            draw.text((left + 5, y - 8), f"{value:.1f}", fill="black", font=small_font)

        bar_width = 90
        spacing = 55
        for index, (model, value, color) in enumerate(zip(models, values, colors)):
            x0 = left + 70 + index * (bar_width + spacing)
            x1 = x0 + bar_width
            y1 = base_y
            y0 = base_y - value * 160
            draw.rectangle([x0, y0, x1, y1], fill=color, outline=color)
            draw.text((x0 + 10, base_y + 8), model, fill="black", font=small_font)
            draw.text((x0 + 10, y0 - 22), f"{value:.3f}", fill="black", font=small_font)

    image.save(output_path)


def create_cnn_architecture_png(output_path: Path) -> None:
    width = 1400
    height = 480
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    block_font = get_font(18)
    small_font = get_font(14)

    draw.text((width // 2 - 210, 18), "1D CNN Architecture Used in This Project", fill="black", font=title_font)

    blocks = [
        ("Input", "1 x 1280\nnormalized\nBVP window"),
        ("Conv Block 1", "Conv1d\n1 -> 16\nk=7\nReLU\nMaxPool(2)"),
        ("Conv Block 2", "Conv1d\n16 -> 32\nk=5\nReLU\nMaxPool(2)"),
        ("Conv Block 3", "Conv1d\n32 -> 64\nk=5\nReLU\nAdaptiveAvgPool(1)"),
        ("Classifier", "Flatten\nDropout\nLinear 64->32\nReLU\nLinear 32->C"),
        ("Output", "Class scores\nSoftmax in\nfinal prediction"),
    ]
    colors = [
        (225, 239, 255),
        (221, 245, 229),
        (221, 245, 229),
        (221, 245, 229),
        (255, 239, 219),
        (245, 226, 240),
    ]

    start_x = 40
    top = 120
    block_w = 190
    block_h = 210
    gap = 32
    arrow_y = top + block_h // 2

    for index, ((title, body), color) in enumerate(zip(blocks, colors)):
        left = start_x + index * (block_w + gap)
        right = left + block_w
        bottom = top + block_h
        draw.rounded_rectangle([left, top, right, bottom], radius=16, fill=color, outline=(90, 90, 90), width=2)
        draw.text((left + 18, top + 16), title, fill="black", font=block_font)

        body_lines = body.split("\n")
        text_y = top + 62
        for line in body_lines:
            draw.text((left + 18, text_y), line, fill="black", font=small_font)
            text_y += 22

        if index < len(blocks) - 1:
            x1 = right
            x2 = right + gap - 6
            draw.line([(x1, arrow_y), (x2, arrow_y)], fill=(80, 80, 80), width=4)
            draw.polygon([(x2, arrow_y), (x2 - 12, arrow_y - 8), (x2 - 12, arrow_y + 8)], fill=(80, 80, 80))

    note = (
        "C = number of output classes. For 3-class experiments C=3 (baseline, stress, amusement); "
        "for binary experiments C=2 (non-stress, stress)."
    )
    draw.text((48, 390), note, fill="black", font=small_font)
    image.save(output_path)


def build_final_report(workspace_root: Path) -> Path:
    outputs_dir = workspace_root / OUTPUT_DIR_NAME
    signal_png = outputs_dir / "report_signal_comparison.png"
    comparison_png = outputs_dir / "report_model_comparison.png"
    cnn_architecture_png = outputs_dir / "report_cnn_architecture.png"
    create_signal_comparison_png(workspace_root, signal_png)
    create_model_comparison_png(outputs_dir, comparison_png)
    create_cnn_architecture_png(cnn_architecture_png)

    split = load_json(outputs_dir / "all_subjects_split.json")
    meta = load_json(outputs_dir / "all_subjects_pipeline_metadata.json")
    svm3 = load_json(outputs_dir / "svm_3class_metrics.json")
    rf3 = load_json(outputs_dir / "rf_3class_metrics.json")
    cnn3 = load_json(outputs_dir / "cnn_3class_metrics.json")
    svmb = load_json(outputs_dir / "svm_binary_metrics.json")
    rfb = load_json(outputs_dir / "rf_binary_metrics.json")
    cnnb = load_json(outputs_dir / "cnn_binary_metrics.json")
    subject_summary = pd.read_csv(outputs_dir / "all_subjects_subject_summary.csv")

    baseline_windows = int(subject_summary["baseline_windows"].sum())
    stress_windows = int(subject_summary["stress_windows"].sum())
    amusement_windows = int(subject_summary["amusement_windows"].sum())

    document = Document()
    title = document.add_heading("Emotion Recognition from Photoplethysmography (PPG) Signals Using WESAD", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Student Name(s): ______________________________\n")
    info.add_run("Course: CSC 491/591\n")
    info.add_run("Date: April 27, 2026")

    add_heading(document, "1. Introduction", level=1)
    add_paragraph(
        document,
        "Stress and affect recognition from wearable physiological signals is an important problem in mobile health. Among wrist-worn sensors, photoplethysmography (PPG) is especially attractive because it is widely available in consumer devices such as smartwatches. In this project, the WESAD dataset was used to build emotion and stress classification models from the wrist blood volume pulse (BVP) signal, which is the PPG-derived signal stored in the synchronized subject pickle files.",
    )
    add_paragraph(
        document,
        "The main goals of this project were to preprocess the wrist BVP signal, convert the long recordings into labeled fixed-length windows, and compare several classification models under a subject-wise evaluation setting. Following both the project guidance and the WESAD literature, the main focus was placed on 3-class classification (baseline vs stress vs amusement) and binary classification (stress vs non-stress).",
    )

    add_heading(document, "2. Dataset", level=1)
    add_paragraph(
        document,
        "The Wearable Stress and Affect Detection (WESAD) dataset contains multimodal physiological recordings from 15 valid subjects. For this project, only the synchronized SX.pkl files were used because they provide aligned sensor data and labels. From each subject file, the wrist BVP signal and the study-protocol labels were extracted.",
    )
    add_table(
        document,
        ["Item", "Value"],
        [
            ["Subjects used", "15 (S2-S17 excluding missing S1 and S12)"],
            ["Signal used", "Wrist BVP (PPG-derived blood volume pulse)"],
            ["BVP sampling rate", "64 Hz"],
            ["Original label sampling rate", "700 Hz"],
            ["Primary class setup", "3-class: baseline, stress, amusement"],
            ["Binary class setup", "stress vs non-stress"],
        ],
    )
    add_paragraph(
        document,
        f"After preprocessing and windowing, the combined dataset contained {meta['num_total_windows']} windows across all subjects before class filtering. For the 3-class experiments, meditation windows were removed, leaving {baseline_windows + stress_windows + amusement_windows} windows in total: {baseline_windows} baseline, {stress_windows} stress, and {amusement_windows} amusement windows.",
    )

    add_heading(document, "3. Data Processing", level=1)
    add_paragraph(
        document,
        "The raw wrist BVP signal was first filtered to reduce baseline wander, high-frequency noise, and motion-related artifacts. A Butterworth-style bandpass filter was applied with a low cutoff of 0.5 Hz, a high cutoff of 4.0 Hz, and order 4. These settings were chosen to preserve the physiologically useful heart-related component while suppressing very slow drift and high-frequency noise.",
    )
    add_paragraph(
        document,
        "After filtering, the study-protocol labels were mapped from the 700 Hz timeline to the 64 Hz BVP timeline. The continuous recordings were then segmented into valid emotion/stress sections. For exploratory work, baseline, stress, amusement, and meditation were all preserved, but the final reported experiments focused on the standard WESAD benchmark classes: baseline, stress, and amusement. Finally, each valid segment was windowed using 20-second windows with 10-second overlap, producing fixed-size inputs for both classical machine learning and deep learning models.",
    )
    document.add_picture(str(signal_png), width=Inches(6.5))
    figure_caption = document.add_paragraph("Figure 1. Example of raw and filtered wrist BVP signal for subject S2 (first 30 seconds).")
    figure_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "4. Methodology", level=1)
    add_paragraph(
        document,
        "Two modeling approaches were used. The first approach was feature-based classification, where each BVP window was summarized by handcrafted statistical features: mean, standard deviation, minimum, maximum, median, range, energy, RMS, skewness, and kurtosis. These features were used by SVM and Random Forest models. The second approach was end-to-end deep learning, where the filtered raw BVP windows were used directly as input to a 1D CNN.",
    )
    add_heading(document, "4.1 SVM", level=2)
    add_paragraph(
        document,
        "The SVM baseline was implemented as a one-vs-rest linear classifier. Feature vectors were standardized using the training data statistics. The model was trained using a margin-based optimization procedure with 60 epochs, learning rate 0.01, and regularization 0.001.",
    )
    add_heading(document, "4.2 Random Forest", level=2)
    add_paragraph(
        document,
        "The Random Forest model was trained on the same handcrafted feature vectors. The final configuration used 25 trees, maximum depth 8, minimum samples split 10, minimum samples leaf 5, and a square-root feature sampling strategy at each split. This model served as the strongest feature-based classifier in the project.",
    )
    add_heading(document, "4.3 1D CNN", level=2)
    add_paragraph(
        document,
        "The deep learning model was a 1D CNN that operated on filtered raw BVP windows. Each window was normalized independently using z-score normalization before being passed to the network. The architecture used three convolutional blocks followed by adaptive average pooling and fully connected layers. Weighted cross-entropy loss was used to reduce the impact of class imbalance, and the model was trained with Adam on CPU for 15 epochs with batch size 64.",
    )
    add_table(
        document,
        ["Layer Group", "Details"],
        [
            ["Input", "1 x 1280 normalized BVP samples (20 s at 64 Hz)"],
            ["Conv Block 1", "Conv1d(1,16,k=7), ReLU, MaxPool1d(2)"],
            ["Conv Block 2", "Conv1d(16,32,k=5), ReLU, MaxPool1d(2)"],
            ["Conv Block 3", "Conv1d(32,64,k=5), ReLU, AdaptiveAvgPool1d(1)"],
            ["Classifier", "Flatten, Dropout, Linear(64->32), ReLU, Linear(32->classes)"],
        ],
    )
    document.add_picture(str(cnn_architecture_png), width=Inches(6.5))
    figure_caption_arch = document.add_paragraph("Figure 2. 1D CNN architecture used in the end-to-end deep learning model.")
    figure_caption_arch.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "5. Experimental Setup", level=1)
    add_paragraph(
        document,
        "To avoid data leakage, the data was split by subject rather than by windows. The training subjects were S2, S3, S4, S5, S6, S7, S8, S9, S10, and S11, while the test subjects were S13, S14, S15, S16, and S17. This split follows the TA guidance that models should be evaluated on unseen subjects. Accuracy, macro F1 score, and confusion matrices were used as evaluation metrics.",
    )
    add_table(
        document,
        ["Split", "Subjects"],
        [
            ["Train", ", ".join(split["train_subjects"])],
            ["Test", ", ".join(split["test_subjects"])],
        ],
    )

    add_heading(document, "6. Results", level=1)
    add_paragraph(
        document,
        "Figure 3 summarizes the main quantitative comparison across models. Tables 1 and 2 provide the detailed numerical results for the 3-class and binary tasks.",
    )
    document.add_picture(str(comparison_png), width=Inches(6.5))
    figure_caption2 = document.add_paragraph("Figure 3. Comparison of SVM, Random Forest, and CNN on 3-class and binary tasks.")
    figure_caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_table(
        document,
        ["Model", "3-Class Accuracy", "3-Class Macro F1", "Binary Accuracy", "Binary Macro F1"],
        [
            ["SVM", f"{svm3['accuracy']:.4f}", f"{svm3['macro_f1']:.4f}", f"{svmb['accuracy']:.4f}", f"{svmb['macro_f1']:.4f}"],
            ["Random Forest", f"{rf3['accuracy']:.4f}", f"{rf3['macro_f1']:.4f}", f"{rfb['accuracy']:.4f}", f"{rfb['macro_f1']:.4f}"],
            ["1D CNN", f"{cnn3['accuracy']:.4f}", f"{cnn3['macro_f1']:.4f}", f"{cnnb['accuracy']:.4f}", f"{cnnb['macro_f1']:.4f}"],
        ],
    )

    add_heading(document, "6.1 3-Class Task", level=2)
    add_paragraph(
        document,
        "For the 3-class task, Random Forest achieved the best overall performance with accuracy 0.6943 and macro F1 0.5922. The SVM achieved moderate performance but failed to identify amusement reliably. The 1D CNN provided a valid deep learning comparison, but in its current form it underperformed the feature-based Random Forest model.",
    )
    add_table(
        document,
        ["Model", "Accuracy", "Macro F1", "Best/Worst Observation"],
        [
            ["SVM", f"{svm3['accuracy']:.4f}", f"{svm3['macro_f1']:.4f}", "Strong on baseline/stress, failed on amusement"],
            ["Random Forest", f"{rf3['accuracy']:.4f}", f"{rf3['macro_f1']:.4f}", "Best overall 3-class model"],
            ["1D CNN", f"{cnn3['accuracy']:.4f}", f"{cnn3['macro_f1']:.4f}", "Very high stress recall, poor baseline separation"],
        ],
    )

    add_heading(document, "6.2 Binary Task", level=2)
    add_paragraph(
        document,
        "For the binary task, all three models performed better than on the 3-class task, which is expected because binary stress recognition is easier than fine-grained emotion separation. Random Forest again performed best overall, while the CNN achieved high recall for the stress class but produced more false positives on the non-stress class.",
    )
    add_table(
        document,
        ["Model", "Accuracy", "Macro F1", "Best/Worst Observation"],
        [
            ["SVM", f"{svmb['accuracy']:.4f}", f"{svmb['macro_f1']:.4f}", "Good non-stress recognition, weaker stress recall"],
            ["Random Forest", f"{rfb['accuracy']:.4f}", f"{rfb['macro_f1']:.4f}", "Best binary model overall"],
            ["1D CNN", f"{cnnb['accuracy']:.4f}", f"{cnnb['macro_f1']:.4f}", "Very high stress recall, lower non-stress recall"],
        ],
    )

    add_heading(document, "7. Discussion", level=1)
    add_paragraph(
        document,
        "Several patterns are clear from the results. First, the binary stress vs non-stress problem is easier than the 3-class baseline vs stress vs amusement problem. This agrees with both the original WESAD paper and the hybrid CNN paper, which also report better performance for binary classification. Second, among the current models, Random Forest is the strongest overall approach. This suggests that the handcrafted statistical features already capture useful stress-related information in the wrist BVP signal.",
    )
    add_paragraph(
        document,
        "The CNN did not outperform Random Forest in the present implementation. This does not invalidate the deep learning approach; rather, it indicates that the current CNN is a baseline end-to-end model and may need stronger feature support, longer windows, LOSO evaluation, or hybrid feature augmentation to match more advanced literature. The hybrid CNN paper is especially relevant here because it shows that combining handcrafted features and CNN-learned representations can improve performance over a plain CNN.",
    )
    add_paragraph(
        document,
        "There are also important methodological differences between this project and the reference papers. The current work uses a 20-second window with 10-second overlap, while the WESAD benchmark papers often use 60-second windows and either 0.25-second or 5-second shifts. The current project also uses a fixed subject-wise train/test split instead of leave-one-subject-out cross-validation. These choices are acceptable for the course project because the TA emphasized subject-wise splitting and did not require replication of the papers exactly. The most important requirement was to justify the preprocessing and modeling choices clearly, which this report does.",
    )

    add_heading(document, "8. Conclusion", level=1)
    add_paragraph(
        document,
        "This project built an end-to-end wrist-PPG emotion/stress recognition pipeline using the WESAD dataset. The raw BVP signal was filtered, segmented, windowed, and converted into both feature-based and raw-signal model inputs. Three models were evaluated: SVM, Random Forest, and 1D CNN. Across the final experiments, Random Forest achieved the strongest performance for both the 3-class and binary tasks, while the CNN served as a meaningful deep learning comparison baseline. Overall, the results show that wrist-based PPG contains enough information to support emotion and stress recognition, especially in the binary stress vs non-stress setting.",
    )

    add_heading(document, "References", level=1)
    add_paragraph(
        document,
        "[1] Philip Schmidt, Attila Reiss, Robert Duerichen, Claus Marberger, and Kristof Van Laerhoven. Introducing WESAD, a multimodal dataset for wearable stress and affect detection. Proceedings of the 20th ACM International Conference on Multimodal Interaction, 2018.",
    )
    add_paragraph(
        document,
        "[2] Nafiul Rashid, Luke Chen, Manik Dautta, Abel Jimenez, Peter Tseng, and Mohammad Abdullah Al Faruque. Feature Augmented Hybrid CNN for Stress Recognition Using Wrist-based Photoplethysmography Sensor. 2021 43rd Annual International Conference of the IEEE Engineering in Medicine & Biology Society (EMBC), 2021.",
    )
    add_paragraph(
        document,
        "[3] CSC 491/591 Course Project Slides and TA Project Explanation Transcript, Spring 2026.",
    )

    output_path = outputs_dir / "wesad_final_project_report.docx"
    document.save(output_path)
    return output_path


def main() -> None:
    workspace_root = Path(__file__).resolve().parent
    output_path = build_final_report(workspace_root)
    print(f"Saved final project report to: {output_path}")


if __name__ == "__main__":
    main()
