from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_two_column_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Value"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def add_metrics_table(document: Document, title: str, metrics: dict) -> None:
    add_heading(document, title, level=2)
    add_two_column_table(
        document,
        [
            ("Accuracy", f"{metrics['accuracy']:.4f}"),
            ("Macro F1", f"{metrics['macro_f1']:.4f}"),
            ("Test Samples", str(metrics["num_test_samples"])),
            ("Class Order", ", ".join(metrics["class_order"])),
        ],
    )

    class_table = document.add_table(rows=1, cols=5)
    class_table.style = "Table Grid"
    headers = class_table.rows[0].cells
    headers[0].text = "Class"
    headers[1].text = "Precision"
    headers[2].text = "Recall"
    headers[3].text = "F1"
    headers[4].text = "Support"

    for class_name, class_metrics in metrics["per_class"].items():
        cells = class_table.add_row().cells
        cells[0].text = class_name
        cells[1].text = f"{class_metrics['precision']:.4f}"
        cells[2].text = f"{class_metrics['recall']:.4f}"
        cells[3].text = f"{class_metrics['f1']:.4f}"
        cells[4].text = str(class_metrics["support"])


def add_model_comparison_table(document: Document, rows: list[tuple[str, str, float, float]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Model"
    headers[1].text = "Task"
    headers[2].text = "Accuracy"
    headers[3].text = "Macro F1"

    for model_name, task_name, accuracy, macro_f1 in rows:
        cells = table.add_row().cells
        cells[0].text = model_name
        cells[1].text = task_name
        cells[2].text = f"{accuracy:.4f}"
        cells[3].text = f"{macro_f1:.4f}"


def add_confusion_matrix_table(document: Document, title: str, csv_path: Path) -> None:
    add_heading(document, title, level=3)
    frame = pd.read_csv(csv_path)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"

    for index, column in enumerate(frame.columns):
        table.rows[0].cells[index].text = str(column)

    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(frame.columns):
            cells[index].text = str(row[column])


def main() -> None:
    workspace_root = Path(__file__).resolve().parent
    outputs_dir = workspace_root / "outputs"

    split = load_json(outputs_dir / "all_subjects_split.json")
    all_subjects_meta = load_json(outputs_dir / "all_subjects_pipeline_metadata.json")
    three_class_metrics = load_json(outputs_dir / "svm_3class_metrics.json")
    binary_metrics = load_json(outputs_dir / "svm_binary_metrics.json")
    rf_three_class_metrics = load_json(outputs_dir / "rf_3class_metrics.json")
    rf_binary_metrics = load_json(outputs_dir / "rf_binary_metrics.json")
    cnn_three_class_metrics = load_json(outputs_dir / "cnn_3class_metrics.json")
    cnn_binary_metrics = load_json(outputs_dir / "cnn_binary_metrics.json")

    document = Document()
    title = document.add_heading("WESAD Project Progress Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Generated from current workspace outputs").italic = True

    add_heading(document, "Current Scope", level=1)
    add_paragraph(
        document,
        "This document summarizes the current preprocessing and SVM baseline results for the WESAD course project.",
    )
    add_bullet_list(
        document,
        [
            "Dataset used: WESAD wrist BVP from SX.pkl files only.",
            "Dataset folders were treated as read-only throughout the workflow.",
            "Signals were filtered with a 4th-order Butterworth-style bandpass response from 0.5 Hz to 4.0 Hz.",
            "Windows were created with 20-second length and 10-second overlap.",
            "Feature set used for SVM baseline: mean, std, min, max, median, range, energy, rms, skewness, kurtosis.",
        ],
    )

    add_heading(document, "Subject Split", level=1)
    add_two_column_table(
        document,
        [
            ("Train Subjects", ", ".join(split["train_subjects"])),
            ("Test Subjects", ", ".join(split["test_subjects"])),
            ("Total Segments", str(all_subjects_meta["num_total_segments"])),
            ("Total Windows", str(all_subjects_meta["num_total_windows"])),
        ],
    )

    add_heading(document, "Label Setups", level=1)
    add_bullet_list(
        document,
        [
            "Research-aligned 3-class setup: baseline, stress, amusement.",
            "Research-aligned binary setup: stress vs non-stress, where non-stress = baseline + amusement.",
        ],
    )

    add_heading(document, "Model Results", level=1)
    add_heading(document, "Model Comparison Summary", level=2)
    add_model_comparison_table(
        document,
        [
            ("SVM", "3-class", three_class_metrics["accuracy"], three_class_metrics["macro_f1"]),
            ("Random Forest", "3-class", rf_three_class_metrics["accuracy"], rf_three_class_metrics["macro_f1"]),
            ("CNN", "3-class", cnn_three_class_metrics["accuracy"], cnn_three_class_metrics["macro_f1"]),
            ("SVM", "binary", binary_metrics["accuracy"], binary_metrics["macro_f1"]),
            ("Random Forest", "binary", rf_binary_metrics["accuracy"], rf_binary_metrics["macro_f1"]),
            ("CNN", "binary", cnn_binary_metrics["accuracy"], cnn_binary_metrics["macro_f1"]),
        ],
    )

    add_metrics_table(document, "3-Class Linear SVM", three_class_metrics)
    add_confusion_matrix_table(
        document,
        "3-Class Confusion Matrix",
        outputs_dir / "svm_3class_confusion_matrix.csv",
    )

    add_metrics_table(document, "Binary Linear SVM", binary_metrics)
    add_confusion_matrix_table(
        document,
        "Binary Confusion Matrix",
        outputs_dir / "svm_binary_confusion_matrix.csv",
    )

    add_metrics_table(document, "3-Class Random Forest", rf_three_class_metrics)
    add_confusion_matrix_table(
        document,
        "3-Class Random Forest Confusion Matrix",
        outputs_dir / "rf_3class_confusion_matrix.csv",
    )

    add_metrics_table(document, "Binary Random Forest", rf_binary_metrics)
    add_confusion_matrix_table(
        document,
        "Binary Random Forest Confusion Matrix",
        outputs_dir / "rf_binary_confusion_matrix.csv",
    )

    add_metrics_table(document, "3-Class CNN", cnn_three_class_metrics)
    add_confusion_matrix_table(
        document,
        "3-Class CNN Confusion Matrix",
        outputs_dir / "cnn_3class_confusion_matrix.csv",
    )

    add_metrics_table(document, "Binary CNN", cnn_binary_metrics)
    add_confusion_matrix_table(
        document,
        "Binary CNN Confusion Matrix",
        outputs_dir / "cnn_binary_confusion_matrix.csv",
    )

    add_heading(document, "Interpretation", level=1)
    add_bullet_list(
        document,
        [
            "The binary stress vs non-stress setup currently performs best among the tested variants.",
            "Random Forest is currently stronger than SVM on both the 3-class task and the binary task.",
            "The CNN gives you the required deep learning comparison, but in the current form it does not beat the feature-based Random Forest baseline.",
            "The 3-class setup is a close match to the original WESAD benchmark and is more informative than the binary task, even though it is harder.",
            "Amusement is still the weakest class overall, but Random Forest improves it substantially compared with the SVM baseline.",
        ],
    )

    add_heading(document, "Key Output Files", level=1)
    add_bullet_list(
        document,
        [
            str(outputs_dir / "all_subjects_window_features.csv"),
            str(outputs_dir / "svm_3class_metrics.json"),
            str(outputs_dir / "svm_binary_metrics.json"),
            str(outputs_dir / "rf_3class_metrics.json"),
            str(outputs_dir / "rf_binary_metrics.json"),
            str(outputs_dir / "cnn_3class_metrics.json"),
            str(outputs_dir / "cnn_binary_metrics.json"),
            str(outputs_dir / "svm_3class_confusion_matrix.svg"),
            str(outputs_dir / "svm_binary_confusion_matrix.svg"),
            str(outputs_dir / "rf_3class_confusion_matrix.svg"),
            str(outputs_dir / "rf_binary_confusion_matrix.svg"),
            str(outputs_dir / "cnn_3class_confusion_matrix.svg"),
            str(outputs_dir / "cnn_binary_confusion_matrix.svg"),
        ],
    )

    add_heading(document, "Next Steps", level=1)
    add_bullet_list(
        document,
        [
            "Use the 3-class and binary setups as the main report experiments.",
            "Keep SVM as the baseline model and Random Forest as the stronger feature-based comparison model.",
            "Keep CNN as the deep learning comparison model, then decide whether to tune it or keep the current result as the first deep baseline.",
        ],
    )

    output_path = outputs_dir / "wesad_progress_report.docx"
    document.save(output_path)
    print(f"Saved report to: {output_path}")


if __name__ == "__main__":
    main()
