from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from generate_final_project_report import (
    OUTPUT_DIR_NAME,
    create_cnn_architecture_png,
    create_signal_comparison_png,
    load_json,
)
from process_s2_pipeline import apply_butterworth_bandpass_fft, load_subject_pickle, map_labels_to_bvp_timeline


def get_font(size: int) -> ImageFont.ImageFont:
    try:
        return ImageFont.truetype("arial.ttf", size)
    except OSError:
        return ImageFont.load_default()


def set_times_new_roman(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.style.font.name = "Times New Roman"


def add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell_run = table.rows[0].cells[index].paragraphs[0].add_run(header)
        cell_run.bold = True
        cell_run.font.name = "Times New Roman"
        cell_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            run = cells[index].paragraphs[0].add_run(value)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def create_loso_model_comparison_png(outputs_dir: Path, output_path: Path) -> None:
    svm_binary = load_json(outputs_dir / "friends_style_binary_svm_metrics.json")
    rf_binary = load_json(outputs_dir / "friends_style_binary_random_forest_metrics.json")
    cnn_binary = load_json(outputs_dir / "friends_style_binary_cnn_metrics.json")
    svm_3class = load_json(outputs_dir / "friends_style_3class_svm_metrics.json")
    rf_3class = load_json(outputs_dir / "friends_style_3class_random_forest_metrics.json")
    cnn_3class = load_json(outputs_dir / "friends_style_3class_cnn_metrics.json")

    models = ["SVM", "Random Forest", "CNN"]
    accuracies_3class = [svm_3class["accuracy"], rf_3class["accuracy"], cnn_3class["accuracy"]]
    macro_f1_3class = [svm_3class["macro_f1"], rf_3class["macro_f1"], cnn_3class["macro_f1"]]
    accuracies_binary = [svm_binary["accuracy"], rf_binary["accuracy"], cnn_binary["accuracy"]]
    macro_f1_binary = [svm_binary["macro_f1"], rf_binary["macro_f1"], cnn_binary["macro_f1"]]
    colors = [(80, 140, 230), (70, 170, 100), (230, 140, 60)]

    width = 1200
    height = 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    text_font = get_font(18)
    small_font = get_font(14)

    draw.text((width // 2 - 150, 20), "LOSO Model Comparison", fill="black", font=title_font)

    panels = [
        ("3-Class Accuracy", accuracies_3class, 60, 110),
        ("3-Class Macro F1", macro_f1_3class, 620, 110),
        ("Binary Accuracy", accuracies_binary, 60, 400),
        ("Binary Macro F1", macro_f1_binary, 620, 400),
    ]

    for title, values, left, top in panels:
        panel_width = 500
        panel_height = 220
        draw.rectangle([left, top, left + panel_width, top + panel_height], outline=(180, 180, 180), width=1)
        draw.text((left + 10, top - 28), title, fill="black", font=text_font)

        base_y = top + panel_height - 30
        draw.line([(left + 40, top + 20), (left + 40, base_y)], fill="black", width=1)
        draw.line([(left + 40, base_y), (left + panel_width - 20, base_y)], fill="black", width=1)

        for tick in range(0, 6):
            value = tick / 5
            y = base_y - value * 160
            draw.line([(left + 40, y), (left + panel_width - 20, y)], fill=(230, 230, 230), width=1)
            draw.text((left + 5, y - 8), f"{value:.1f}", fill="black", font=small_font)

        bar_width = 90
        spacing = 55
        for index, (model, value, color) in enumerate(zip(models, values, colors)):
            x0 = left + 70 + index * (bar_width + spacing)
            x1 = x0 + bar_width
            y1 = base_y
            y0 = base_y - value * 160
            draw.rectangle([x0, y0, x1, y1], fill=color, outline=color)
            draw.text((x0 + 14, base_y + 8), model, fill="black", font=small_font)
            draw.text((x0 + 14, y0 - 22), f"{value:.3f}", fill="black", font=small_font)

    image.save(output_path)


def create_processed_segments_png(workspace_root: Path, output_path: Path) -> None:
    subject_id = "S2"
    data = load_subject_pickle(subject_id, workspace_root)
    bvp = np.asarray(data["signal"]["wrist"]["BVP"]).flatten()
    labels_700hz = np.asarray(data["label"])
    filtered = apply_butterworth_bandpass_fft(
        signal=bvp,
        sampling_rate=64,
        low_cutoff_hz=0.7,
        high_cutoff_hz=3.7,
        order=3,
    )
    labels_64hz = map_labels_to_bvp_timeline(
        labels_700hz=labels_700hz,
        bvp_length=len(filtered),
        bvp_sampling_rate=64,
        label_sampling_rate=700,
    )

    class_ids = [("baseline", 1), ("stress", 2), ("amusement", 3)]
    window_samples = 1800
    segments: list[tuple[str, np.ndarray]] = []
    for label_name, label_id in class_ids:
        indices = np.where(labels_64hz == label_id)[0]
        if len(indices) == 0:
            continue
        start = int(indices[0])
        end = min(start + window_samples, len(filtered))
        segment = filtered[start:end]
        segments.append((label_name, segment))

    width = 1250
    height = 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    axis_font = get_font(18)
    tick_font = get_font(14)

    draw.text((width // 2 - 210, 20), "Processed PPG Segments (Filtered BVP)", fill="black", font=title_font)
    colors = {"baseline": (70, 130, 210), "stress": (230, 140, 60), "amusement": (90, 190, 110)}

    margin_left = 100
    margin_right = 40
    margin_top = 90
    margin_bottom = 40
    panel_gap = 34
    panel_height = (height - margin_top - margin_bottom - 2 * panel_gap) // 3
    plot_width = width - margin_left - margin_right

    for panel_index, (label_name, signal) in enumerate(segments):
        top = margin_top + panel_index * (panel_height + panel_gap)
        y_min = float(np.min(signal))
        y_max = float(np.max(signal))
        if y_max == y_min:
            y_max = y_min + 1.0

        def sx(x_value: float) -> float:
            return margin_left + x_value / max(1, len(signal) - 1) * plot_width

        def sy(y_value: float) -> float:
            return top + panel_height - ((y_value - y_min) / (y_max - y_min)) * panel_height

        draw.rectangle([margin_left, top, margin_left + plot_width, top + panel_height], outline=(180, 180, 180), width=1)
        draw.text((margin_left, top - 28), f"Emotion: {label_name}", fill="black", font=axis_font)

        for tick in range(0, 6):
            x_value = tick / 5
            x = margin_left + x_value * plot_width
            draw.line([(x, top), (x, top + panel_height)], fill=(235, 235, 235), width=1)
        for tick in range(0, 5):
            y_value = y_min + (y_max - y_min) * tick / 4
            y = sy(y_value)
            draw.line([(margin_left, y), (margin_left + plot_width, y)], fill=(235, 235, 235), width=1)
            draw.text((20, y - 8), f"{y_value:.0f}", fill="black", font=tick_font)

        points = [(sx(index), sy(float(value))) for index, value in enumerate(signal)]
        draw.line(points, fill=colors[label_name], width=2)

    draw.text((width // 2 - 90, height - 28), "Time (samples)", fill="black", font=axis_font)
    image.save(output_path)


def create_report_confusion_matrix_png(matrix: np.ndarray, class_labels: list[str], title: str, output_path: Path) -> None:
    cell_size = 105
    left_margin = 160
    top_margin = 100
    width = left_margin + cell_size * len(class_labels) + 60
    height = top_margin + cell_size * len(class_labels) + 70
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(24)
    label_font = get_font(16)
    value_font = get_font(18)
    max_value = int(matrix.max()) if matrix.size else 1

    title_width = draw.textbbox((0, 0), title, font=title_font)[2]
    draw.text(((width - title_width) / 2, 24), title, fill="black", font=title_font)
    pred_width = draw.textbbox((0, 0), "Predicted label", font=label_font)[2]
    draw.text(((width - pred_width) / 2, 58), "Predicted label", fill="black", font=label_font)
    draw.text((20, top_margin - 28), "True label", fill="black", font=label_font)

    def cell_color(value: int) -> tuple[int, int, int]:
        intensity = value / max_value if max_value else 0.0
        blue = int(255 - 120 * intensity)
        red = int(245 - 110 * intensity)
        green = int(248 - 120 * intensity)
        return red, green, blue

    for column_index, label in enumerate(class_labels):
        x = left_margin + column_index * cell_size + 18
        draw.text((x, top_margin - 28), label, fill="black", font=label_font)
    for row_index, label in enumerate(class_labels):
        y = top_margin + row_index * cell_size + 42
        draw.text((20, y), label, fill="black", font=label_font)

    for row_index, row_values in enumerate(matrix):
        for column_index, value in enumerate(row_values):
            x0 = left_margin + column_index * cell_size
            y0 = top_margin + row_index * cell_size
            x1 = x0 + cell_size
            y1 = y0 + cell_size
            draw.rectangle([x0, y0, x1, y1], fill=cell_color(int(value)), outline=(90, 90, 90), width=2)
            draw.text((x0 + 35, y0 + 40), str(int(value)), fill="black", font=value_font)

    image.save(output_path)


def create_model_pipeline_png(title: str, blocks: list[tuple[str, str]], output_path: Path) -> None:
    width = 1380
    height = 430
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(28)
    block_font = get_font(18)
    small_font = get_font(14)

    title_width = draw.textbbox((0, 0), title, font=title_font)[2]
    draw.text(((width - title_width) / 2, 18), title, fill="black", font=title_font)

    colors = [
        (225, 239, 255),
        (221, 245, 229),
        (255, 245, 220),
        (245, 226, 240),
        (232, 240, 246),
        (252, 236, 221),
    ]

    start_x = 40
    top = 120
    block_w = 190
    block_h = 190
    gap = 28
    arrow_y = top + block_h // 2

    for index, (block_title, body) in enumerate(blocks):
        left = start_x + index * (block_w + gap)
        right = left + block_w
        bottom = top + block_h
        fill = colors[index % len(colors)]
        draw.rounded_rectangle([left, top, right, bottom], radius=16, fill=fill, outline=(90, 90, 90), width=2)
        draw.text((left + 16, top + 16), block_title, fill="black", font=block_font)
        text_y = top + 54
        for line in body.split("\n"):
            draw.text((left + 16, text_y), line, fill="black", font=small_font)
            text_y += 22
        if index < len(blocks) - 1:
            x1 = right
            x2 = right + gap - 6
            draw.line([(x1, arrow_y), (x2, arrow_y)], fill=(80, 80, 80), width=4)
            draw.polygon([(x2, arrow_y), (x2 - 12, arrow_y - 8), (x2 - 12, arrow_y + 8)], fill=(80, 80, 80))

    image.save(output_path)


def metrics_rows(metrics: dict[str, object]) -> list[list[str]]:
    per_class = metrics["per_class"]
    rows = [["Overall", f"{metrics['accuracy']:.4f}", f"{metrics['macro_f1']:.4f}", "-", "-"]]
    for label in metrics["class_order"]:
        rows.append(
            [
                label,
                f"{per_class[label]['precision']:.4f}",
                f"{per_class[label]['recall']:.4f}",
                f"{per_class[label]['f1']:.4f}",
                str(per_class[label]["support"]),
            ]
        )
    return rows


def summary_row_3class(model_name: str, metrics: dict[str, object]) -> list[str]:
    return [
        model_name,
        f"{metrics['accuracy']:.3f}",
        f"{metrics['macro_f1']:.3f}",
        f"{metrics['per_class']['baseline']['f1']:.3f}",
        f"{metrics['per_class']['stress']['f1']:.3f}",
        f"{metrics['per_class']['amusement']['f1']:.3f}",
    ]


def summary_row_binary(model_name: str, metrics: dict[str, object]) -> list[str]:
    return [
        model_name,
        f"{metrics['accuracy']:.3f}",
        f"{metrics['macro_f1']:.3f}",
        f"{metrics['per_class']['non-stress']['f1']:.3f}",
        f"{metrics['per_class']['stress']['f1']:.3f}",
    ]


def save_document_with_fallback(document: Document, output_path: Path) -> Path:
    try:
        document.save(output_path)
        return output_path
    except PermissionError:
        for suffix in ["_updated", "_v2", "_v3", "_v4", "_final"]:
            fallback_path = output_path.with_name(f"{output_path.stem}{suffix}{output_path.suffix}")
            try:
                document.save(fallback_path)
                return fallback_path
            except PermissionError:
                continue
        raise


def build_final_loso_report(workspace_root: Path) -> Path:
    outputs_dir = workspace_root / OUTPUT_DIR_NAME
    signal_png = outputs_dir / "report_signal_comparison.png"
    processed_segments_png = outputs_dir / "report_processed_segments.png"
    svm_pipeline_png = outputs_dir / "report_svm_pipeline.png"
    rf_pipeline_png = outputs_dir / "report_random_forest_pipeline.png"
    cnn_architecture_png = outputs_dir / "report_cnn_architecture.png"
    comparison_png = outputs_dir / "report_loso_model_comparison.png"
    binary_svm_cm = outputs_dir / "report_binary_svm_confusion.png"
    binary_rf_cm = outputs_dir / "report_binary_random_forest_confusion.png"
    binary_cnn_cm = outputs_dir / "report_binary_cnn_confusion.png"
    class3_svm_cm = outputs_dir / "report_3class_svm_confusion.png"
    class3_rf_cm = outputs_dir / "report_3class_random_forest_confusion.png"
    class3_cnn_cm = outputs_dir / "report_3class_cnn_confusion.png"

    create_signal_comparison_png(workspace_root, signal_png)
    create_processed_segments_png(workspace_root, processed_segments_png)
    create_model_pipeline_png(
        "SVM Pipeline Used in This Project",
        [
            ("Input", "60 s filtered\nBVP window"),
            ("Feature Extraction", "24 handcrafted\nPPG / HRV-style\nfeatures"),
            ("Standardization", "mean = 0\nstd = 1 using\ntraining fold"),
            ("Classifier", "One-vs-rest\nlinear SVM"),
            ("Output", "baseline /\nstress /\namusement\nor binary class"),
        ],
        svm_pipeline_png,
    )
    create_model_pipeline_png(
        "Random Forest Pipeline Used in This Project",
        [
            ("Input", "60 s filtered\nBVP window"),
            ("Feature Extraction", "24 handcrafted\nPPG / HRV-style\nfeatures"),
            ("Tree Ensemble", "25 trees\nmax depth = 8\nmin split = 10\nmin leaf = 5"),
            ("Voting", "aggregate\npredictions from\nall trees"),
            ("Output", "baseline /\nstress /\namusement\nor binary class"),
        ],
        rf_pipeline_png,
    )
    create_cnn_architecture_png(cnn_architecture_png)
    create_loso_model_comparison_png(outputs_dir, comparison_png)

    svm_binary = load_json(outputs_dir / "friends_style_binary_svm_metrics.json")
    rf_binary = load_json(outputs_dir / "friends_style_binary_random_forest_metrics.json")
    cnn_binary = load_json(outputs_dir / "friends_style_binary_cnn_metrics.json")
    svm_3class = load_json(outputs_dir / "friends_style_3class_svm_metrics.json")
    rf_3class = load_json(outputs_dir / "friends_style_3class_random_forest_metrics.json")
    cnn_3class = load_json(outputs_dir / "friends_style_3class_cnn_metrics.json")
    binary_summary = load_json(outputs_dir / "friends_style_binary_summary.json")
    class3_summary = load_json(outputs_dir / "friends_style_3class_summary.json")
    binary_frame = pd.read_csv(outputs_dir / "friends_style_binary_dataset_features.csv")
    class3_frame = pd.read_csv(outputs_dir / "friends_style_3class_dataset_features.csv")

    binary_counts = binary_frame["target_label"].value_counts().to_dict()
    class3_counts = class3_frame["target_label"].value_counts().to_dict()
    subject_count = int(binary_frame["subject"].nunique())

    create_report_confusion_matrix_png(np.asarray(svm_binary["confusion_matrix"]), svm_binary["class_order"], "SVM Confusion Matrix", binary_svm_cm)
    create_report_confusion_matrix_png(np.asarray(rf_binary["confusion_matrix"]), rf_binary["class_order"], "Random Forest Confusion Matrix", binary_rf_cm)
    create_report_confusion_matrix_png(np.asarray(cnn_binary["confusion_matrix"]), cnn_binary["class_order"], "CNN Confusion Matrix", binary_cnn_cm)
    create_report_confusion_matrix_png(np.asarray(svm_3class["confusion_matrix"]), svm_3class["class_order"], "SVM Confusion Matrix", class3_svm_cm)
    create_report_confusion_matrix_png(np.asarray(rf_3class["confusion_matrix"]), rf_3class["class_order"], "Random Forest Confusion Matrix", class3_rf_cm)
    create_report_confusion_matrix_png(np.asarray(cnn_3class["confusion_matrix"]), cnn_3class["class_order"], "CNN Confusion Matrix", class3_cnn_cm)

    document = Document()
    set_times_new_roman(document)

    title = document.add_heading("Final Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Course 591")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(document, "Name : - Yuvraj Singh Bhatia (ybhatia2)")
    add_paragraph(document, "Date: April 27, 2026")

    add_heading(document, "1. Overview", level=1)
    add_paragraph(
        document,
        "This report presents emotion and stress recognition experiments on the WESAD dataset using wrist-based photoplethysmography (PPG). The objective was to evaluate whether the wrist blood volume pulse signal can distinguish baseline, stress, and amusement states, and whether stress can be separated from non-stress under a subject-independent setting.",
    )
    add_paragraph(
        document,
        "To make the evaluation realistic, Leave-One-Subject-Out (LOSO) validation was used. In this setting, each subject is held out once for testing while the remaining subjects are used for training. Three models were evaluated: Support Vector Machine (SVM), Random Forest, and a one-dimensional Convolutional Neural Network (CNN). Results are reported for both 3-class classification and binary classification.",
    )

    add_heading(document, "2. Data Processing", level=1)
    add_paragraph(
        document,
        "The synchronized SX.pkl files from WESAD were used. Only the wrist BVP signal and the protocol labels were retained. The signal was filtered using a Butterworth-style bandpass filter with a low cutoff of 0.7 Hz, a high cutoff of 3.7 Hz, and order 3. This step reduces baseline drift and high-frequency noise while preserving the frequency range that carries useful cardiac information.",
    )
    add_paragraph(
        document,
        "After filtering, each subject signal was normalized with subject-wise z-score normalization. Fixed-length windows of 60 seconds were extracted with a step size of 5 seconds, and windows containing mixed labels were discarded. For the binary setting, baseline and amusement windows were merged into a non-stress class, while stress windows remained unchanged. For the 3-class setting, baseline, stress, and amusement were kept as separate labels.",
    )
    add_table(
        document,
        ["Item", "Value"],
        [
            ["Subjects used", str(subject_count)],
            ["Signal", "Wrist BVP"],
            ["Filter", "0.7-3.7 Hz, order 3"],
            ["Window size", "60 seconds"],
            ["Step size", "5 seconds"],
            ["Evaluation", "LOSO"],
            ["3-class windows", str(class3_summary["dataset"]["num_windows"])],
            ["Binary windows", str(binary_summary["dataset"]["num_windows"])],
        ],
    )
    add_paragraph(
        document,
        f"3-class label counts: baseline {class3_counts.get('baseline', 0)}, stress {class3_counts.get('stress', 0)}, amusement {class3_counts.get('amusement', 0)}. Binary label counts: non-stress {binary_counts.get('non-stress', 0)}, stress {binary_counts.get('stress', 0)}.",
    )
    document.add_picture(str(signal_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Figure 1. Example of raw and filtered wrist BVP signal used in preprocessing.")
    document.add_picture(str(processed_segments_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Figure 2. Example processed PPG segments for baseline, stress, and amusement after filtering.")

    add_heading(document, "3. Methodology", level=1)
    add_paragraph(
        document,
        "Three classification models were implemented in order to compare feature-based learning with end-to-end deep learning. The first two models, SVM and Random Forest, used the same handcrafted PPG feature set. The CNN used normalized raw windows directly as input.",
    )

    add_paragraph(document, "Support Vector Machine (SVM)", bold=True)
    add_paragraph(
        document,
        "The SVM model was implemented as a one-vs-rest linear classifier. Input features were standardized before training. This model provides a strong margin-based baseline and is well suited to smaller physiological datasets.",
    )
    document.add_picture(str(svm_pipeline_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Figure 3. SVM processing pipeline used in the LOSO experiments.")

    add_paragraph(document, "Random Forest", bold=True)
    add_paragraph(
        document,
        "The Random Forest model was trained on the same handcrafted features. The forest used 25 trees, a maximum depth of 8, a minimum split size of 10, a minimum leaf size of 5, and square-root feature sampling at each split. This model is useful because it can capture nonlinear feature interactions without requiring extensive feature scaling.",
    )
    document.add_picture(str(rf_pipeline_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Figure 4. Random Forest processing pipeline used in the LOSO experiments.")

    add_paragraph(document, "CNN", bold=True)
    add_paragraph(
        document,
        "The CNN model was designed as a one-dimensional convolutional network over 60-second normalized BVP windows. It used stacked convolutional layers, ReLU activations, pooling, and a small fully connected classifier. For the binary LOSO setting, the final CNN was trained for 10 epochs per fold. For the 3-class LOSO setting, the final CNN was also trained for 10 epochs per fold. Weighted loss was used to reduce the effect of class imbalance.",
    )
    document.add_picture(str(cnn_architecture_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Figure 5. CNN architecture used for the LOSO experiments.")

    add_heading(document, "4. Experiments and Analysis", level=1)
    add_paragraph(
        document,
        "Performance was measured using accuracy, macro F1 score, and confusion matrices. Accuracy indicates overall correctness, while macro F1 emphasizes balance across classes. This is especially important in physiological classification tasks where class behavior is not equally easy to learn.",
    )
    add_paragraph(document, "3-Class Evaluation Metrics", bold=True)
    add_table(
        document,
        ["Model", "Accuracy", "Macro F1", "F1-Baseline", "F1-Stress", "F1-Amusement"],
        [
            summary_row_3class("SVM", svm_3class),
            summary_row_3class("Random Forest", rf_3class),
            summary_row_3class("CNN", cnn_3class),
        ],
    )
    add_paragraph(document, "Binary Evaluation Metrics", bold=True)
    add_table(
        document,
        ["Model", "Accuracy", "Macro F1", "F1-Non-Stress", "F1-Stress"],
        [
            summary_row_binary("SVM", svm_binary),
            summary_row_binary("Random Forest", rf_binary),
            summary_row_binary("CNN", cnn_binary),
        ],
    )
    document.add_picture(str(comparison_png), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Figure 6. Comparison of accuracy and macro F1 across models.")

    add_paragraph(document, "Binary Classification Metrics", bold=True)
    add_paragraph(document, "SVM", bold=True)
    add_table(document, ["Class", "Precision/Accuracy", "Recall/Macro F1", "F1", "Support"], metrics_rows(svm_binary))
    add_paragraph(document, "Random Forest", bold=True)
    add_table(document, ["Class", "Precision/Accuracy", "Recall/Macro F1", "F1", "Support"], metrics_rows(rf_binary))
    add_paragraph(document, "CNN", bold=True)
    add_table(document, ["Class", "Precision/Accuracy", "Recall/Macro F1", "F1", "Support"], metrics_rows(cnn_binary))

    add_paragraph(document, "3-Class Classification Metrics", bold=True)
    add_paragraph(document, "SVM", bold=True)
    add_table(document, ["Class", "Precision/Accuracy", "Recall/Macro F1", "F1", "Support"], metrics_rows(svm_3class))
    add_paragraph(document, "Random Forest", bold=True)
    add_table(document, ["Class", "Precision/Accuracy", "Recall/Macro F1", "F1", "Support"], metrics_rows(rf_3class))
    add_paragraph(document, "CNN", bold=True)
    add_table(document, ["Class", "Precision/Accuracy", "Recall/Macro F1", "F1", "Support"], metrics_rows(cnn_3class))

    add_paragraph(
        document,
        f"In the binary LOSO setting, SVM achieved the strongest overall result with accuracy {svm_binary['accuracy']:.4f} and macro F1 {svm_binary['macro_f1']:.4f}. The CNN improved substantially after training for 10 epochs and reached accuracy {cnn_binary['accuracy']:.4f} with macro F1 {cnn_binary['macro_f1']:.4f}, which made it highly competitive with the classical baselines.",
    )
    add_paragraph(
        document,
        f"In the 3-class LOSO setting, SVM gave the highest accuracy at {svm_3class['accuracy']:.4f}, while CNN achieved the highest macro F1 at {cnn_3class['macro_f1']:.4f}. This suggests that SVM made the most correct predictions overall, whereas CNN produced the most balanced performance across baseline, stress, and amusement.",
    )

    add_paragraph(document, "Confusion Matrix Analysis", bold=True)
    add_paragraph(
        document,
        "The confusion matrices provide a more detailed view of model behavior than accuracy and macro F1 alone. In the binary setting, the main source of error is confusion between stress and non-stress windows. In the 3-class setting, amusement is the most challenging class, while baseline and stress are generally recognized more consistently.",
    )
    add_paragraph(document, "Binary Confusion Matrices", bold=True)
    add_paragraph(document, "SVM Confusion Matrix", bold=True)
    document.add_picture(str(binary_svm_cm), width=Inches(4.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Random Forest Confusion Matrix", bold=True)
    document.add_picture(str(binary_rf_cm), width=Inches(4.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "CNN Confusion Matrix", bold=True)
    document.add_picture(str(binary_cnn_cm), width=Inches(4.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "3-Class Confusion Matrices", bold=True)
    add_paragraph(document, "SVM Confusion Matrix", bold=True)
    document.add_picture(str(class3_svm_cm), width=Inches(4.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "Random Forest Confusion Matrix", bold=True)
    document.add_picture(str(class3_rf_cm), width=Inches(4.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "CNN Confusion Matrix", bold=True)
    document.add_picture(str(class3_cnn_cm), width=Inches(4.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "5. Conclusion", level=1)
    add_paragraph(
        document,
        "This study showed that wrist-based PPG from WESAD can support both binary stress recognition and 3-class emotion recognition under a strict LOSO evaluation protocol. The preprocessing pipeline, longer windows, and subject-wise normalization produced stable inputs for all three models.",
    )
    add_paragraph(
        document,
        "Among the evaluated models, SVM was the strongest binary classifier and also produced the highest 3-class accuracy. The CNN benefited significantly from longer training, and after 10 epochs it achieved the highest 3-class macro F1, indicating improved class balance. Overall, the results show that both classical machine learning and deep learning are viable for PPG-based affect recognition, with the best model depending on whether overall accuracy or balanced class performance is prioritized.",
    )

    output_path = outputs_dir / "Final_Report_LOSO.docx"
    return save_document_with_fallback(document, output_path)


def main() -> None:
    workspace_root = Path(__file__).resolve().parent
    output_path = build_final_loso_report(workspace_root)
    print(f"Saved LOSO Word report to: {output_path}")


if __name__ == "__main__":
    main()
